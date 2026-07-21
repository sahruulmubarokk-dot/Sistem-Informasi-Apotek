import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# 1. Inisialisasi Dokumen & Set Margin Kaku (Left 4cm, Top/Right/Bottom 3cm)
doc = docx.Document()

# Helper to add dynamic page number to a run (using XML PAGE field)
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def set_section_page_number_format(section, format_str='romanLower'):
    secPr = section._sectPr
    pgNumType = secPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        secPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), format_str)

def set_section_page_number_start(section, start_val=1):
    secPr = section._sectPr
    pgNumType = secPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        secPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start_val))

# SECTION 1: Cover & Front Matter (LEMBAR PENGESAHAN, KATA PENGANTAR, DAFTAR ISI)
sec1 = doc.sections[0]
sec1.top_margin = Inches(1.18)      # 3 cm
sec1.bottom_margin = Inches(1.18)   # 3 cm
sec1.left_margin = Inches(1.57)     # 4 cm
sec1.right_margin = Inches(1.18)    # 3 cm
sec1.different_first_page_header_footer = True

# Set page numbering to lowercase Roman numerals (i, ii, iii...)
set_section_page_number_format(sec1, 'romanLower')

# Add footer to Section 1
f_foot1 = sec1.footer
p_foot1 = f_foot1.paragraphs[0]
p_foot1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_foot1 = p_foot1.add_run()
run_foot1.font.name = 'Times New Roman'
run_foot1.font.size = Pt(11)
add_page_number(run_foot1)

# Helper Pemformatan Teks Bersih (Clean-Code Spacing & Times New Roman)
def add_custom_p(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

# Custom multiline helper to split text by newline and add them as separate paragraphs
def add_custom_p_multiline(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, bold=False, size=12):
    lines = text.split('\n')
    added = []
    for line in lines:
        if line.strip():  # Skip empty lines
            added.append(add_custom_p(line.strip(), align=align, space_after=space_after, bold=bold, size=size))
    return added

def add_heading_1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

def add_heading_2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

def add_image_to_doc(image_path, caption_text):
    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if image_path in ["diagram_konteks.png", "dfd_level0.png", "erd_database.png"]:
            p_img.add_run().add_picture(image_path, width=Inches(5.0), height=Inches(2.96))
        else:
            p_img.add_run().add_picture(image_path, width=Inches(5.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(12)
        run = p_cap.add_run(caption_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.italic = True
    else:
        add_custom_p(f"\n[Gambar '{image_path}' Tidak Ditemukan]\n", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10)

def format_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = bold

def add_toc_entry(title, page_str, is_sub=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25) if is_sub else Inches(0)
    
    # 8.5 - 1.57 - 1.18 = 5.75 inches printable width. Set tab stop at 5.75 right.
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(5.75), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    run = p.add_run(f"{title}\t{page_str}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    if not is_sub:
        run.bold = True
    return p

# =========================================================================
# [A. COVER & HALAMAN AWAL] -> Section 1
# =========================================================================
add_custom_p("LAPORAN AKHIR PROJECT-BASED LEARNING (PBL)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
add_custom_p("RANCANG BANGUN SISTEM INFORMASI APOTEK BERBASIS WEB\n(APOTEKSIM) DENGAN ARSITEKTUR MODERN FULL-STACK\nDAN INTEGRASI SUPABASE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)

# Logo Kampus
logo_path = "UNITOMO.jpg"
if os.path.exists(logo_path):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(18)
    p_logo.paragraph_format.space_after = Pt(18)
    p_logo.add_run().add_picture(logo_path, width=Inches(2.36), height=Inches(2.36))
else:
    add_custom_p("\n[Tempatkan File Logo Universitas 'UNITOMO.jpg' di Sini]\n", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10)

add_custom_p("Disusun Oleh Kelompok 4:\n", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_custom_p("Fahri Adis Al Hafni – NIM: 202511420105 (Role: Fullstack Developer & DB Designer)", align=WD_ALIGN_PARAGRAPH.CENTER)
add_custom_p("Sahrul Mubarok – NIM: 202511420107 (Role: Backend & Cloud Infrastructure)", align=WD_ALIGN_PARAGRAPH.CENTER)
add_custom_p("Ach Tohir – NIM: 202511420048 (Role: Frontend UI/UX & QA Tester)", align=WD_ALIGN_PARAGRAPH.CENTER)

add_custom_p("\nPROGRAM STUDI TEKNIK INFORMATIKA\nFAKULTAS TEKNIK\nUNIVERSITAS DR. SOETOMO\nSURABAYA\n2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

doc.add_page_break()

# Lembar Pengesahan
add_heading_1("LEMBAR PENGESAHAN")
add_custom_p("Laporan akhir Project-Based Learning (PBL) dengan judul 'Rancang Bangun Sistem Informasi Apotek Berbasis Web (ApotekSim) dengan Arsitektur Modern Full-Stack dan Integrasi Supabase' ini telah diperiksa, disetujui, dan disahkan sebagai salah satu syarat akademis dalam menyelesaikan mata kuliah Project-Based Learning pada Program Studi Teknik Informatika, Fakultas Teknik, Universitas Dr. Soetomo Surabaya.")

# Tabel Lembar Pengesahan 1 (Menyetujui: Dosen & Kaprodi)
table_sign1 = doc.add_table(rows=1, cols=2)
table_sign1.style = 'Normal Table'
table_sign1.autofit = False

# Dosen Pengampu
cell_dosen = table_sign1.cell(0, 0)
format_cell(cell_dosen, "Mengetahui,\nDosen Pengampu PBL\n\n\n\n(_________________________)")

# Kaprodi
cell_kaprodi = table_sign1.cell(0, 1)
format_cell(cell_kaprodi, "Surabaya, 12 Juli 2026\nKetua Program Studi\n\n\n\n(_________________________)")

doc.add_paragraph() # Jeda

# Tabel Lembar Pengesahan 2 (Dibuat Oleh Kelompok 4)
add_custom_p("Dibuat Oleh Kelompok 4:", bold=True)
table_sign2 = doc.add_table(rows=1, cols=3)
table_sign2.style = 'Normal Table'
table_sign2.autofit = False

# Ketua
format_cell(table_sign2.cell(0, 0), "Ketua / Fullstack\n\n\n\nFahri Adis Al Hafni\nNIM. 202511420105")
# Anggota 1
format_cell(table_sign2.cell(0, 1), "Anggota / Backend\n\n\n\nSahrul Mubarok\nNIM. 202511420107")
# Anggota 2
format_cell(table_sign2.cell(0, 2), "Anggota / Frontend\n\n\n\nAch Tohir\nNIM. 202511420048")

doc.add_page_break()

# Kata Pengantar
add_heading_1("KATA PENGANTAR")
add_custom_p("Puji syukur kehadirat Allah SWT atas rahmat-Nya sehingga laporan akhir PBL ini dapat diselesaikan dengan baik dan tepat waktu. Proyek rancang bangun sistem informasi manajemen apotek terintegrasi (ApotekSim) ini berfokus pada efisiensi transaksi, otomatisasi inventaris berbasis metode FEFO (First Expired, First Out), serta keamanan data transaksional menggunakan arsitektur modern full-stack berbasis cloud (Next.js & Supabase).")
add_custom_p("Terima kasih kami sampaikan kepada Dekan Fakultas Teknik, Ketua Program Studi Teknik Informatika Universitas Dr. Soetomo, Dosen Pengampu mata kuliah, serta rekan-rekan Kelompok 4 atas kolaborasi intensifnya selama pengerjaan proyek ini.")

doc.add_page_break()

# =========================================================================
# [DAFTAR ISI]
# =========================================================================
add_heading_1("DAFTAR ISI")
add_toc_entry("LEMBAR PENGESAHAN", "ii")
add_toc_entry("KATA PENGANTAR", "iii")
add_toc_entry("DAFTAR ISI", "iv")

add_toc_entry("BAB I PENDAHULUAN", "1")
add_toc_entry("1.1 Latar Belakang Masalah", "1", is_sub=True)
add_toc_entry("1.2 Rumusan Masalah", "2", is_sub=True)
add_toc_entry("1.3 Batasan Masalah", "2", is_sub=True)
add_toc_entry("1.4 Tujuan Proyek", "2", is_sub=True)
add_toc_entry("1.5 Manfaat Proyek", "2", is_sub=True)

add_toc_entry("BAB II LANDASAN TEORI", "3")
add_toc_entry("2.1 Sistem Informasi Manajemen (SIM) Apotek", "3", is_sub=True)
add_toc_entry("2.2 Manajemen Inventaris Metode FEFO", "3", is_sub=True)
add_toc_entry("2.3 Arsitektur Modern Full-Stack (Next.js & TypeScript)", "4", is_sub=True)
add_toc_entry("2.4 Cloud Backend-as-a-Service (Supabase & PostgreSQL)", "4", is_sub=True)

add_toc_entry("BAB III ANALISIS DAN PERANCANGAN SISTEM", "5")
add_toc_entry("3.1 Diagram Konteks", "5", is_sub=True)
add_toc_entry("3.2 Data Flow Diagram (DFD) Level 0", "6", is_sub=True)
add_toc_entry("3.3 Entity Relationship Diagram (ERD)", "7", is_sub=True)
add_toc_entry("3.4 Perancangan Arsitektur Basis Data Relasional", "8", is_sub=True)
add_toc_entry("3.5 Alur Kerja Bisnis Sistem (End-to-End Workflow)", "9", is_sub=True)

add_toc_entry("BAB IV IMPLEMENTASI DAN PENGUJIAN", "10")
add_toc_entry("4.1 Deskripsi Lingkungan Implementasi", "10", is_sub=True)
add_toc_entry("4.2 Keamanan Row Level Security (RLS) Supabase", "10", is_sub=True)
add_toc_entry("4.3 Hasil Tampilan Program", "11", is_sub=True)
add_toc_entry("4.4 Skenario Pengujian Sistem (Black-Box Testing)", "13", is_sub=True)

add_toc_entry("BAB V PENUTUP", "14")
add_toc_entry("5.1 Kesimpulan", "14", is_sub=True)
add_toc_entry("5.2 Saran Pengembangan Selanjutnya", "14", is_sub=True)

add_toc_entry("DAFTAR PUSTAKA", "15")

# =========================================================================
# SECTION 2: Body (BAB I to DAFTAR PUSTAKA)
# =========================================================================
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
sec2.top_margin = Inches(1.18)      # 3 cm
sec2.bottom_margin = Inches(1.18)   # 3 cm
sec2.left_margin = Inches(1.57)     # 4 cm
sec2.right_margin = Inches(1.18)    # 3 cm
sec2.different_first_page_header_footer = False

# Set page numbering format to Arabic (1, 2, 3...) and restart at 1
set_section_page_number_format(sec2, 'arabic')
set_section_page_number_start(sec2, 1)

# Unlink footer from Section 1 and configure footer for Section 2
f_foot2 = sec2.footer
f_foot2.is_linked_to_previous = False
p_foot2 = f_foot2.paragraphs[0]
p_foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_foot2 = p_foot2.add_run()
run_foot2.font.name = 'Times New Roman'
run_foot2.font.size = Pt(11)
add_page_number(run_foot2)

# =========================================================================
# [B. BAB I PENDAHULUAN] -> Section 2 Starts
# =========================================================================
add_heading_1("BAB I PENDAHULUAN")
add_heading_2("1.1 Latar Belakang Masalah")
add_custom_p("Apotek merupakan salah satu pilar krusial dalam rantai pasok pelayanan kesehatan masyarakat. Operasional harian apotek menuntut ketelitian tinggi, mulai dari pelacakan ketersediaan obat, pengelolaan tanggal kedaluwarsa, pencatatan resep dokter secara legal, hingga pencatatan transaksi kasir harian (Point of Sales). Berdasarkan observasi lapangan pada sejumlah apotek skala menengah ke bawah, sebagian besar manajemen operasional masih mengandalkan pencatatan manual berbasis buku atau bantuan aplikasi lembar sebar (spreadsheet) lokal yang tidak terintegrasi.")
add_custom_p("Kondisi manajemen yang belum terdigitalisasi secara andal ini menimbulkan berbagai risiko operasional yang fatal, di antaranya kerugian finansial akibat obat kedaluwarsa yang tidak terdeteksi sejak dini karena ketiadaan pencatatan berbasis batch. Selain itu, human error dalam stok sering terjadi akibat proses transaksi kasir yang tidak langsung memotong jumlah inventaris secara real-time. Yang tidak kalah berbahaya adalah pelanggaran regulasi obat keras; penjualan obat kategori keras, narkotika, atau psikotropika tanpa pengawasan data resep dokter yang tervalidasi dapat berujung pada sanksi administratif hingga pencabutan izin operasional apotek.")
add_custom_p("Seiring perkembangan teknologi informasi, transformasi digital di ranah operasional apotek menjadi sebuah keharusan. Oleh karena itu, diperlukan sebuah solusi sistem informasi berbasis web modern yang responsif, mengutamakan validitas data transaksional, serta aman secara menyeluruh. Melalui program Project-Based Learning (PBL) ini, Kelompok 4 merancang dan membangun ApotekSim, sebuah Sistem Informasi Apotek berbasis web dengan arsitektur modern full-stack menggunakan Next.js (TypeScript) dan didukung oleh infrastruktur cloud Supabase (PostgreSQL) di sisi backend. Sistem ini mengutamakan manajemen batch obat terotomatisasi, sistem antrean FEFO, sinkronisasi data real-time, dan pengamanan akses berbasis peran (RBAC) yang ketat.")

add_heading_2("1.2 Rumusan Masalah")
add_custom_p_multiline(
    "1. Bagaimana merancang arsitektur basis data relasional yang ternormalisasi untuk sistem manajemen apotek yang melibatkan entitas Obat, Supplier, dan Pembelian?\n"
    "2. Bagaimana mengimplementasikan sistem antrean pengeluaran barang berbasis First Expired, First Out (FEFO) secara otomatis pada saat proses transaksi kasir (Point of Sales) berlangsung?\n"
    "3. Bagaimana membangun sistem informasi apotek dengan mematuhi prinsip Clean Code dan Separation of Concerns (SoC) menggunakan integrasi kerangka kerja Next.js dan Supabase?\n"
    "4. Bagaimana mengamankan transaksi data finansial dan operasional apotek dari manipulasi siber sisi klien menggunakan kebijakan keamanan database yang kokoh?"
)

add_heading_2("1.3 Batasan Masalah")
add_custom_p_multiline(
    "1. Sistem mencakup pengelolaan data master obat, supplier, pencatatan transaksi pembelian (restock) berbasis batch, pencatatan transaksi kasir (POS), skrining resep dokter, dan dasbor analitik visual sederhana.\n"
    "2. Basis data dan layanan infrastruktur server sepenuhnya menggunakan ekosistem Supabase BaaS (PostgreSQL).\n"
    "3. Sisi frontend menggunakan Next.js, TypeScript, Tailwind CSS, dan komponen Shadcn/ui.\n"
    "4. Sistem diasumsikan berjalan secara online-first agar sinkronisasi data real-time dapat berfungsi dengan baik."
)

add_heading_2("1.4 Tujuan Proyek")
add_custom_p_multiline(
    "1. Menghasilkan rancangan arsitektur basis data relasional yang skalabel dan ternormalisasi guna menghindari redundansi data obat, supplier, dan transaksi.\n"
    "2. Mengembangkan antarmuka pengguna kasir digital (POS) yang cepat, intuitif, dan responsif guna menekan waktu tunggu transaksi.\n"
    "3. Mengotomatisasikan pelacakan tanggal kedaluwarsa serta batas minimum stok obat guna mempermudah proses pengambilan keputusan pengadaan obat kembali.\n"
    "4. Mengimplementasikan kebijakan keamanan tingkat baris data (Row Level Security) pada Supabase guna memastikan hak akses data yang presisi antara Admin, Apoteker, dan Kasir."
)

add_heading_2("1.5 Manfaat Proyek")
add_custom_p("Membantu pemilik apotek mengoptimalkan efisiensi kerja staf, meminimalkan kerugian finansial akibat obat kedaluwarsa lewat sistem FEFO otomatis, mempermudah kasir mengeliminasi proses pengecekan fisik secara manual, serta menjadi referensis akademis penerapan arsitektur web modern full-stack (Serverless / Cloud-native) dalam memecahkan kasus bisnis dunia nyata.")

doc.add_page_break()

# =========================================================================
# [C. BAB II LANDASAN TEORI]
# =========================================================================
add_heading_1("BAB II LANDASAN TEORI")
add_heading_2("2.1 Sistem Informasi Manajemen (SIM) Apotek")
add_custom_p("Sistem Informasi Manajemen (SIM) berfungsi mengintegrasikan seluruh alur informasi barang masuk (pengadaan), barang keluar (penjualan kasir), sirkulasi keuangan, serta dokumentasi administratif medis secara konsisten pada setiap titik operasional guna menghindari terjadinya selisih data stok maupun kesalahan kalkulasi kasir.")

add_heading_2("2.2 Manajemen Inventaris Metode FEFO (First Expired, First Out)")
add_custom_p("Dalam industri farmasi, metode FEFO menetapkan bahwa barang yang memiliki tanggal kedaluwarsa paling dekat harus dikeluarkan atau dijual terlebih dahulu. Formulasi alokasi stok berdasarkan antrean FEFO didefinisikan sebagai berikut:")
add_custom_p("Misalkan suatu jenis obat O memiliki beberapa batch persediaan B = {b1, b2, ..., bn}, di mana setiap batch memiliki kuantitas q_i dan tanggal kedaluwarsa E(b_i) yang diurutkan berdasarkan relasi:")
add_custom_p("E(b1) <= E(b2) <= ... <= E(bn)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_custom_p("Ketika terjadi permintaan transaksi penjualan sejumlah unit Q, sistem harus mengalokasikan pengurangan stok secara berurutan mulai dari b1. Untuk setiap batch b_i dalam urutan tersebut, jumlah pengurangan d_i ditentukan oleh formula:")
add_custom_p("d_i = min(Q_sisa, q_i)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_custom_p("Di mana Q_sisa diperbarui pada setiap iterasi: Q_sisa^(i) = Q_sisa^(i-1) - d_i dengan Q_sisa^(0) = Q. Proses ini terus berjalan hingga Q_sisa = 0.")

add_heading_2("2.3 Arsitektur Modern Full-Stack (Next.js & TypeScript)")
add_custom_p("Next.js sebagai kerangka kerja React modern mengintegrasikan Server-Side Rendering (SSR) dan Static Site Generation (SSG) dengan Hydration di sisi klien. Penggunaan TypeScript menambahkan lapisan keamanan statis (static type-safety) di atas JavaScript sehingga mencegah runtime errors yang dapat mengacaukan perhitungan matematis total omset penjualan.")

add_heading_2("2.4 Cloud Backend-as-a-Service (Supabase & PostgreSQL)")
add_custom_p("Supabase merupakan platform BaaS sumber terbuka yang ditenagai oleh mesin basis data PostgreSQL. Fitur utama yang dimanfaatkan meliputi Supabase Auth (manajemen JWT), Supabase Realtime (WebSockets sync), dan PostgreSQL Row Level Security (RLS) untuk perlindungan mutasi data di sisi server.")

doc.add_page_break()

# =========================================================================
# [D. BAB III ANALISIS DAN PERANCANGAN SISTEM]
# =========================================================================
add_heading_1("BAB III ANALISIS DAN PERANCANGAN SISTEM")
add_heading_2("3.1 Diagram Konteks")
add_custom_p("Diagram konteks di bawah ini menggambarkan batasan sistem informasi manajemen ApotekSim dengan entitas luar (Admin, Apoteker, Kasir, Supplier):")
add_image_to_doc("diagram_konteks.png", "Gambar 3.1 Diagram Konteks Sistem ApotekSim")

add_heading_2("3.2 Data Flow Diagram (DFD) Level 0")
add_custom_p("DFD Level 0 memecah sistem menjadi sub-proses utama: Katalogisasi Master Data, Manajemen Pasokan (Restock Batch), Penjualan Kasir (POS) Berbasis FEFO, dan Pelaporan Real-Time:")
add_image_to_doc("dfd_level0.png", "Gambar 3.2 Data Flow Diagram (DFD) Level 0")

add_heading_2("3.3 Entity Relationship Diagram (ERD)")
add_custom_p("ERD di bawah ini memetakan hubungan entitas logis logistik dan transaksional pada database ApotekSim:")
add_image_to_doc("erd_database.png", "Gambar 3.3 Entity Relationship Diagram (ERD) Relasi Database ApotekSim")

add_heading_2("3.4 Perancangan Arsitektur Basis Data Relasional")
add_custom_p("Berikut adalah spesifikasi teknis komponen database relasional utama yang diterapkan:")

# Tabel 1: Supplier
add_custom_p("1. Tabel Master: Supplier", bold=True)
t_sup = doc.add_table(rows=6, cols=4)
t_sup.style = 'Light Shading Accent 1'
headers_sup = ["Nama Field", "Tipe Data", "Constraint", "Keterangan"]
for col_idx, header in enumerate(headers_sup):
    format_cell(t_sup.cell(0, col_idx), header, bold=True)
data_sup = [
    ["ID_Supplier", "UUID", "PK, Default: gen_random_uuid()", "ID Unik Supplier"],
    ["Nama_Supplier", "VARCHAR(150)", "NOT NULL", "Nama Vendor Pemasok"],
    ["No_Telepon", "VARCHAR(20)", "NOT NULL", "Nomor Kontak Aktif"],
    ["Alamat", "TEXT", "NOT NULL", "Alamat Fisik Kantor Vendor"],
    ["Created_At", "TIMESTAMP", "Default: now()", "Waktu Registrasi Data"]
]
for row_idx, row_data in enumerate(data_sup):
    for col_idx, cell_data in enumerate(row_data):
        format_cell(t_sup.cell(row_idx + 1, col_idx), cell_data)

doc.add_paragraph() # Spacing

# Tabel 2: Obat
add_custom_p("2. Tabel Master: Obat", bold=True)
t_obt = doc.add_table(rows=7, cols=4)
t_obt.style = 'Light Shading Accent 1'
headers_obt = ["Nama Field", "Tipe Data", "Constraint", "Keterangan"]
for col_idx, header in enumerate(headers_obt):
    format_cell(t_obt.cell(0, col_idx), header, bold=True)
data_obt = [
    ["ID_Obat", "UUID", "PK, Default: gen_random_uuid()", "ID Unik Obat"],
    ["Nama_Obat", "VARCHAR(150)", "NOT NULL, UNIQUE", "Nama Dagang Obat"],
    ["Kategori_Obat", "VARCHAR(50)", "NOT NULL", "Pengelompokan Medis"],
    ["Harga_Jual", "DECIMAL(12, 2)", "NOT NULL", "Harga Jual ke Pelanggan"],
    ["Batas_Minimum_Stok", "INT", "Default: 10", "Threshold Alert Stok Tipis"],
    ["Stok_Total", "INT", "Default: 0", "Kuantitas Agregat Seluruh Batch"]
]
for row_idx, row_data in enumerate(data_obt):
    for col_idx, cell_data in enumerate(row_data):
        format_cell(t_obt.cell(row_idx + 1, col_idx), cell_data)

doc.add_paragraph() # Spacing

# Tabel 3: Pembelian_Obat
add_custom_p("3. Tabel Transaksi Induk: Pembelian_Obat (Restock)", bold=True)
t_pem = doc.add_table(rows=7, cols=4)
t_pem.style = 'Light Shading Accent 1'
headers_pem = ["Nama Field", "Tipe Data", "Constraint", "Keterangan"]
for col_idx, header in enumerate(headers_pem):
    format_cell(t_pem.cell(0, col_idx), header, bold=True)
data_pem = [
    ["ID_Pembelian", "UUID", "PK, Default: gen_random_uuid()", "ID Unik Faktur Masuk"],
    ["No_Faktur", "VARCHAR(100)", "NOT NULL, UNIQUE", "Nomor Faktur dari Supplier"],
    ["ID_Supplier", "UUID", "FK (Supplier.ID_Supplier)", "Referensi Supplier Pemasok"],
    ["Tanggal_Pembelian", "DATE", "Default: CURRENT_DATE", "Tanggal Penerimaan Barang"],
    ["Total_Pengeluaran", "DECIMAL(15, 2)", "NOT NULL", "Total Finansial Faktur"],
    ["Admin_Penginput", "UUID", "FK (Auth.users)", "User yang Melakukan Input"]
]
for row_idx, row_data in enumerate(data_pem):
    for col_idx, cell_data in enumerate(row_data):
        format_cell(t_pem.cell(row_idx + 1, col_idx), cell_data)

doc.add_paragraph() # Spacing

# Tabel 4: Detail_Pembelian_Obat
add_custom_p("4. Tabel Transaksi Detail: Detail_Pembelian_Obat (Batch Inventory)", bold=True)
t_det = doc.add_table(rows=9, cols=4)
t_det.style = 'Light Shading Accent 1'
headers_det = ["Nama Field", "Tipe Data", "Constraint", "Keterangan"]
for col_idx, header in enumerate(headers_det):
    format_cell(t_det.cell(0, col_idx), header, bold=True)
data_det = [
    ["ID_Detail_Pembelian", "UUID", "PK, Default: gen_random_uuid()", "ID Unik Detail Item"],
    ["ID_Pembelian", "UUID", "FK (Pembelian_Obat.ID_Pembelian)", "Relasi Induk Faktur"],
    ["ID_Obat", "UUID", "FK (Obat.ID_Obat)", "Relasi Produk Obat"],
    ["Nomor_Batch", "VARCHAR(50)", "NOT NULL", "Kode Batch Produksi Pabrik"],
    ["Jumlah_Masuk", "INT", "NOT NULL, > 0", "Kuantitas Masuk"],
    ["Stok_Sisa_Batch", "INT", "NOT NULL", "Sisa Kuantitas Batch (FEFO Target)"],
    ["Harga_Beli_Satuan", "DECIMAL(12, 2)", "NOT NULL", "Harga Pokok Pembelian"],
    ["Tanggal_Kedaluwarsa", "DATE", "NOT NULL", "Acuan Utama Algoritma FEFO"]
]
for row_idx, row_data in enumerate(data_det):
    for col_idx, cell_data in enumerate(row_data):
        format_cell(t_det.cell(row_idx + 1, col_idx), cell_data)

add_heading_2("3.5 Alur Kerja Bisnis Sistem (End-to-End Workflow)")
add_custom_p("Sistem diawali dengan pendaftaran data master supplier dan katalog produk obat oleh Administrator. Saat obat baru dipasok, Admin memasukkan faktur pembelian detail dengan nomor batch dan tanggal kedaluwarsa spesifik per item. Data ini masuk ke detail pembelian dan secara otomatis memperbarui Stok_Total pada tabel Obat via trigger database.")
add_custom_p("Pada sisi kasir, menu Point of Sales memanggil katalog obat aktif yang diurutkan sekuensial berdasarkan tanggal kedaluwarsa terdekat (FEFO Order). Apabila kasir melakukan checkout, sistem memicu fungsi transaksi server-side secara atomik guna mengurangi Stok_Sisa_Batch pada batch terdekat serta Stok_Total pada katalog utama obat, meminimalisir kegagalan sinkronisasi data.")

doc.add_page_break()

# =========================================================================
# [E. BAB IV IMPLEMENTASI DAN PENGUJIAN]
# =========================================================================
add_heading_1("BAB IV IMPLEMENTASI DAN PENGUJIAN")
add_heading_2("4.1 Deskripsi Lingkungan Implementasi")
add_custom_p("ApotekSim dijalankan dengan Next.js 14, TypeScript, Tailwind CSS, serta library antarmuka Shadcn/ui. Backend terintegrasi penuh dengan ekosistem database cloud Supabase PostgreSQL.")

add_heading_2("4.2 Keamanan Row Level Security (RLS) Supabase")
add_custom_p("Untuk melindungi data dari manipulasi langsung sisi klien, kebijakan Row Level Security diaktifkan pada PostgreSQL. Hak akses menulis (INSERT, UPDATE, DELETE) dikunci ketat khusus untuk pengguna terautentikasi dengan JSON Web Token (JWT) yang memiliki metadata role = 'admin'. Sedangkan staf kasir dan apoteker hanya diberikan otorisasi membaca (SELECT).")

add_heading_2("4.3 Hasil Tampilan Program")
add_custom_p("Berikut adalah hasil visualisasi dari antarmuka dinamis sistem informasi ApotekSim:")
add_image_to_doc("ui_katalog_obat.png", "Gambar 4.1 Visualisasi Antrean Batch FEFO Aktif")
add_image_to_doc("ui_pos_kasir.png", "Gambar 4.2 Antarmuka Point of Sales (POS) dengan Skrining Resep")
add_image_to_doc("ui_dashboard_admin.png", "Gambar 4.3 Panel Kontrol Dashboard Manajemen dan Log Audit Trail")
add_image_to_doc("ui_tabel_restock.png", "Gambar 4.4 Form Faktur Pembelian (Restock) dan Log Transaksi Multi-Item")

add_heading_2("4.4 Skenario Pengujian Sistem (Black-Box Testing)")
add_custom_p("Sistem diuji menggunakan metode Black-Box Testing untuk memvalidasi kesesuaian operasional sistem:")

t_test = doc.add_table(rows=5, cols=4)
t_test.style = 'Light Shading Accent 1'
headers_test = ["ID Uji", "Skenario Pengujian", "Hasil yang Diharapkan", "Status"]
for col_idx, header in enumerate(headers_test):
    format_cell(t_test.cell(0, col_idx), header, bold=True)
data_test = [
    ["TC-01", "Akses Dasbor Admin via Akun Kasir", "Sistem mendeteksi role JWT kasir dan melakukan redirect proteksi (Access Denied).", "BERHASIL"],
    ["TC-02", "Transaksi Kasir Kategori Obat Keras", "Sistem mengunci tombol bayar dan menampilkan form skrining resep dokter wajib isi.", "BERHASIL"],
    ["TC-03", "Checkout Multi-Batch FEFO Otomatis", "Sistem memotong stok Batch A (ED terdekat) hingga habis, lalu melanjutkan sisa potongan ke Batch B.", "BERHASIL"],
    ["TC-04", "Konsolidasi Alert Stok Minimum", "Ketika stok total jatuh di bawah threshold minimal, indikator visual di dasbor menyala merah secara real-time.", "BERHASIL"]
]
for row_idx, row_data in enumerate(data_test):
    for col_idx, cell_data in enumerate(row_data):
        format_cell(t_test.cell(row_idx + 1, col_idx), cell_data)

doc.add_page_break()

# =========================================================================
# [F. BAB V PENUTUP & DAFTAR PUSTAKA]
# =========================================================================
add_heading_1("BAB V PENUTUP")
add_heading_2("5.1 Kesimpulan")
add_custom_p_multiline(
    "1. Purwarupa Sistem Informasi Apotek (ApotekSim) berhasil dibangun dengan mengimplementasikan arsitektur modern full-stack menggunakan Next.js dan Supabase cloud database.\n"
    "2. Perancangan basis data relasional yang dinormalisasi dengan memisahkan tabel master obat statis dan tabel mutasi detail batch berhasil meminimalisir redundansi data.\n"
    "3. Otomatisasi pengeluaran barang berbasis metode FEFO berhasil berjalan secara transaksional di tingkat database server-side, menekan risiko human error dalam distribusi obat kedaluwarsa.\n"
    "4. Kebijakan Row Level Security terbukti andal menyaring hak akses mutasi data secara langsung berdasarkan validasi identitas JWT pengguna."
)

add_heading_2("5.2 Saran Pengembangan Selanjutnya")
add_custom_p_multiline(
    "1. Disarankan bagi pengembangan selanjutnya untuk mengintegrasikan modul kemampuan Offline-First menggunakan Progressive Web App (PWA) agar operasional kasir tidak terganggu saat internet putus.\n"
    "2. Mengintegrasikan API interaksi obat medis otomatis guna mendeteksi potensi bahaya kombinasi resep dokter secara digital."
)

doc.add_page_break()
add_heading_1("DAFTAR PUSTAKA")
add_custom_p("1. Connolly, T., & Begg, C. (2015). Database Systems: A Practical Approach to Design, Implementation, and Management (6th ed.). Boston: Pearson Education.")
add_custom_p("2. Fowler, M. (2018). Refactoring: Improving the Design of Existing Code (2nd ed.). Boston: Addison-Wesley.")
add_custom_p("3. Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). New York: McGraw-Hill Education.")
add_custom_p("4. Supabase. (2026). Supabase Documentation: Database Policies and Row Level Security. Diakses dari https://supabase.com/docs.")
add_custom_p("5. Vercel. (2026). Next.js Documentation: Server Actions and Client Components. Diakses dari https://nextjs.org/docs.")

# 3. Save Hasil Laporan Akhir PBL (.docx)
base_dir = os.path.dirname(os.path.abspath(__file__))
reports_dir = os.path.abspath(os.path.join(base_dir, "..", "reports"))
os.makedirs(reports_dir, exist_ok=True)
target_path = os.path.join(reports_dir, "laporan_akhir_pbl_apotek.docx")

try:
    doc.save(target_path)
    print(f"Sukses total men-generate file '{target_path}' dengan spasi 1.5 baris dan margin kaku akademik!")
except PermissionError:
    fallback_path = os.path.join(reports_dir, "laporan_akhir_pbl_apotek_new.docx")
    doc.save(fallback_path)
    print(f"Sukses men-generate file '{fallback_path}' (file utama sedang dikunci/dibuka oleh program lain)!")

